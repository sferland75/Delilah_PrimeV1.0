"""
Delilah Prime - Web Application

This module provides a web interface for the Delilah Prime application.
"""

import os
import sys
import json
import glob
from pathlib import Path
import datetime
from flask import Flask, render_template, request, redirect, url_for, flash, send_from_directory, jsonify, session, Response
from werkzeug.utils import secure_filename
from dotenv import load_dotenv
import logging
import uuid
import time
from collections import defaultdict
from queue import Queue, Empty
from threading import Lock

# Make sure we can import from sibling directories
sys.path.append(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

# Import Delilah modules
from src.deidentifier.deidentifier import Deidentifier
from src.api.claude_api import ClaudeAPIClient
from src.processor.document_processor import DocumentProcessor
from src.processor.report_generator import ReportGenerator

# Word document support
try:
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_EXPORT = True
except ImportError:
    DOCX_EXPORT = False
    print("Warning: python-docx not installed, Word export disabled")

# Load environment variables
load_dotenv()

# Configure logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(name)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

# Configuration
OFFLINE_MODE = os.environ.get('OFFLINE_MODE', 'false').lower() == 'true'  # Set to true to skip API calls
ADMIN_MODE = os.environ.get('ADMIN_MODE', 'false').lower() == 'true'  # Enable admin features like Prompt Lab

# Initialize Flask app
app = Flask(__name__, template_folder=os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), 'templates'))
app.secret_key = os.environ.get('ENCRYPTION_KEY', 'delilah_prime_secure_key')
app.config['SESSION_TYPE'] = 'filesystem'
app.config['PERMANENT_SESSION_LIFETIME'] = datetime.timedelta(days=1)

# Set up file paths
BASE_PATH = Path(__file__).parent.parent
INPUT_PATH = BASE_PATH / "input"
OUTPUT_PATH = BASE_PATH / "output"
TEMPLATES_PATH = BASE_PATH / "templates"
PROMPTS_PATH = BASE_PATH / "prompts"
CUSTOM_TEMPLATES_PATH = BASE_PATH / "custom_templates"

# Ensure directories exist
os.makedirs(INPUT_PATH, exist_ok=True)
os.makedirs(OUTPUT_PATH, exist_ok=True)
os.makedirs(TEMPLATES_PATH, exist_ok=True)
os.makedirs(PROMPTS_PATH, exist_ok=True)
os.makedirs(CUSTOM_TEMPLATES_PATH, exist_ok=True)

# Load configuration
CONFIG_PATH = BASE_PATH / "config.json"
with open(CONFIG_PATH, 'r') as f:
    CONFIG = json.load(f)

# Initialize components
deidentifier = Deidentifier(ref_table_path=Path.home() / Path(CONFIG["paths"]["reference_tables"]))
document_processor = DocumentProcessor(CONFIG_PATH)
report_generator = ReportGenerator(CONFIG_PATH)
claude_api = ClaudeAPIClient(CONFIG_PATH)

# Allowed file extensions
ALLOWED_EXTENSIONS = {'txt', 'md', 'docx', 'pdf'}
REPORT_SECTIONS = CONFIG['report_sections']

# Global event queue for streaming activity
event_queue = Queue()
event_queue_lock = Lock()

def log_event(message):
    """Add an event to the queue for streaming"""
    with event_queue_lock:
        event_queue.put(f"[{datetime.datetime.now().strftime('%H:%M:%S')}] {message}")
        # Also log to the console
        app.logger.info(message)

def allowed_file(filename):
    """Check if a file has an allowed extension."""
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

@app.route('/')
def index():
    """Render the index page with input and output files."""
    # Get input files
    input_files = [file.name for file in INPUT_PATH.glob('*.*') 
                   if file.suffix.lower()[1:] in ALLOWED_EXTENSIONS]
    
    # Get output files
    output_files = [file.name for file in OUTPUT_PATH.glob('*.*')]
    
    # Check for reference tables
    ref_tables = list(Path.home().glob(f"{CONFIG['paths']['reference_tables']}/ref_table_*.json"))
    
    # Check for draft report
    draft_exists = (OUTPUT_PATH / "draft_report.txt").exists()
    
    # Custom prompts
    if ADMIN_MODE:
        custom_prompts = [file.name for file in PROMPTS_PATH.glob('*.txt')]
    else:
        custom_prompts = []
    
    # Available templates
    custom_templates = [file.stem for file in CUSTOM_TEMPLATES_PATH.glob('*.json')]
    
    return render_template('index.html', 
                          input_files=input_files, 
                          output_files=output_files,
                          api_available=claude_api.is_available(),
                          has_ref_table=len(ref_tables) > 0,
                          offline_mode=OFFLINE_MODE,
                          admin_mode=ADMIN_MODE,
                          custom_prompts=custom_prompts,
                          custom_templates=custom_templates,
                          draft_exists=draft_exists,
                          report_sections=REPORT_SECTIONS)

@app.route('/upload', methods=['POST'])
def upload_file():
    """Handle file upload."""
    if 'file' not in request.files:
        flash('No file part')
        return redirect(request.url)
    
    file = request.files['file']
    
    if file.filename == '':
        flash('No selected file')
        return redirect(request.url)
    
    if file and allowed_file(file.filename):
        filename = secure_filename(file.filename)
        file.save(os.path.join(INPUT_PATH, filename))
        flash(f'File {filename} uploaded successfully')
        log_event(f"File uploaded: {filename}")
        return redirect(url_for('index'))
    
    flash('Invalid file type')
    return redirect(url_for('index'))

@app.route('/process')
def process_files():
    """Process uploaded files and generate report."""
    try:
        log_event("Beginning document processing...")
        
        # Get input files
        input_files = [os.path.join(INPUT_PATH, f) for f in os.listdir(INPUT_PATH) 
                    if os.path.isfile(os.path.join(INPUT_PATH, f))]
        
        if not input_files:
            flash('No input files found')
            return redirect(url_for('index'))
            
        # Get selected template and prompt if provided
        template_name = request.args.get('template', 'default')
        prompt_name = request.args.get('prompt', '')
        
        log_event(f"Using template: {template_name}, prompt: {prompt_name}")
        
        # Prepare template and prompt
        report_template = None
        prompt_template = None
        
        if template_name != 'default':
            template_path = os.path.join(TEMPLATES_PATH, f"{template_name}.txt")
            if os.path.exists(template_path):
                with open(template_path, 'r', encoding='utf-8') as f:
                    report_template = f.read()
                    log_event(f"Loaded template: {template_name}")
        
        if prompt_name:
            prompt_path = os.path.join(TEMPLATES_PATH, 'prompts', f"{prompt_name}.txt")
            if os.path.exists(prompt_path):
                with open(prompt_path, 'r', encoding='utf-8') as f:
                    prompt_template = f.read()
                    log_event(f"Loaded prompt: {prompt_name}")
        
        # Process all documents to organize content by section
        log_event("Organizing content from files...")
        organized_content = document_processor.organize_content(input_files)
        
        # Save organized content for reference
        with open(os.path.join(OUTPUT_PATH, 'organized_content.json'), 'w', encoding='utf-8') as f:
            json.dump(organized_content, f, indent=2)
            
        log_event("Saved organized content for reference")
        
        # Process each section with AI enhancement if available
        log_event("Starting report section processing...")
        reference_table = deidentifier.get_reference_table()
        
        try:
            # Process all sections at once to leverage our concurrency protection
            enhanced_content = document_processor.process_documents(input_files, prompt_template, reference_table)
            
            # Generate the report
            log_event("Generating final report...")
            report_content = report_generator.generate_report(enhanced_content, template=report_template)
            
            # Save draft report
            draft_path = os.path.join(OUTPUT_PATH, 'draft_report.txt')
            with open(draft_path, 'w', encoding='utf-8') as f:
                f.write(report_content)
                
            log_event("Saved draft report")
            
            # Save reference table if available
            ref_table_path = os.path.join(OUTPUT_PATH, 'reference_table.json')
            if reference_table:
                with open(ref_table_path, 'w', encoding='utf-8') as f:
                    json.dump(reference_table, f, indent=2)
                    log_event("Saved reference table")
            
            # Store report in session for refinement
            session['report_content'] = enhanced_content
            session['has_ref_table'] = os.path.exists(ref_table_path)
            
            flash('Report generated successfully')
            log_event("Report generation complete - ready for review")
        except Exception as e:
            log_event(f"Error in report processing: {str(e)}")
            # Fallback to organized content if processing fails
            log_event("Using organized content as fallback...")
            
            # Save the organized content directly
            with open(os.path.join(OUTPUT_PATH, 'enhanced_content.json'), 'w', encoding='utf-8') as f:
                json.dump(organized_content, f, indent=2)
                
            # Generate a simple report from the organized content
            simple_report = "# Generated Report\n\n"
            for section_name, content in organized_content.items():
                if content:
                    formatted_section = section_name.replace('_', ' ').title()
                    simple_report += f"## {formatted_section}\n\n{content}\n\n"
            
            # Save simple report
            draft_path = os.path.join(OUTPUT_PATH, 'draft_report.txt')
            with open(draft_path, 'w', encoding='utf-8') as f:
                f.write(simple_report)
                
            log_event("Saved simple report as fallback")
            
            # Store report in session for refinement
            session['report_content'] = organized_content
            session['has_ref_table'] = False
            
            flash('Report generated with basic content (API enhancement failed)')
            log_event("Basic report generation complete")
            
        return redirect(url_for('index'))
    except Exception as e:
        error_msg = f"Error processing files: {str(e)}"
        log_event(f"ERROR: {error_msg}")
        flash(error_msg)
        return redirect(url_for('index'))

@app.route('/preview_report')
def preview_report():
    """Preview the draft report with options to refine sections."""
    draft_path = OUTPUT_PATH / "draft_report.txt"
    
    if not draft_path.exists():
        flash('No draft report available. Please process your files first.')
        return redirect(url_for('index'))
    
    with open(draft_path, 'r', encoding='utf-8') as f:
        content = f.read()
    
    # Load organized content if available
    organized_path = OUTPUT_PATH / "organized_content.json"
    if organized_path.exists():
        with open(organized_path, 'r', encoding='utf-8') as f:
            original_content = json.load(f)
    else:
        original_content = {}
    
    # Extract sections from draft report
    sections = {}
    current_section = "preamble"
    section_text = []
    
    for line in content.split('\n'):
        if line.startswith('# ') and not line.startswith('# Clinical Assessment Report'):
            # Save previous section
            if section_text:
                sections[current_section] = '\n'.join(section_text)
            
            # Start new section
            current_section = line[2:].lower().replace(' ', '_')
            section_text = [line]
        else:
            section_text.append(line)
    
    # Save last section
    if section_text:
        sections[current_section] = '\n'.join(section_text)
    
    return render_template('preview_report.html', 
                          content=content,
                          sections=sections,
                          original_content=original_content,
                          report_sections=REPORT_SECTIONS)

@app.route('/refine_section', methods=['POST'])
def refine_section():
    """Refine a specific section with Claude API."""
    section = request.form.get('section')
    content = request.form.get('content')
    custom_prompt = request.form.get('custom_prompt')
    
    if not section or not content:
        flash('Missing section or content')
        return redirect(url_for('preview_report'))
    
    try:
        if claude_api.is_available() and not OFFLINE_MODE:
            # Use custom prompt if provided
            if custom_prompt:
                enhanced_content = claude_api.generate_custom_narrative(section, content, custom_prompt)
            else:
                enhanced_content = claude_api.generate_narrative(section, content)
            
            # Update draft report
            draft_path = OUTPUT_PATH / "draft_report.txt"
            with open(draft_path, 'r', encoding='utf-8') as f:
                report_content = f.read()
            
            # Replace section in report
            section_title = section.replace('_', ' ').title()
            section_pattern = f"# {section_title}\n"
            section_start = report_content.find(section_pattern)
            
            if section_start >= 0:
                next_section_start = report_content.find("# ", section_start + len(section_pattern))
                
                if next_section_start >= 0:
                    # Replace section until next section
                    report_content = report_content[:section_start] + enhanced_content + report_content[next_section_start:]
                else:
                    # Replace until end of document
                    report_content = report_content[:section_start] + enhanced_content
            
            # Save updated draft
            with open(draft_path, 'w', encoding='utf-8') as f:
                f.write(report_content)
            
            flash(f'Section "{section_title}" refined successfully')
        else:
            flash('Claude API not available or offline mode enabled')
        
        return redirect(url_for('preview_report'))
    
    except Exception as e:
        flash(f'Error refining section: {str(e)}')
        return redirect(url_for('preview_report'))

@app.route('/manual_edit', methods=['POST'])
def manual_edit():
    """Manually edit the draft report content."""
    content = request.form.get('content')
    
    if not content:
        flash('No content provided')
        return redirect(url_for('preview_report'))
    
    try:
        # Save updated draft
        draft_path = OUTPUT_PATH / "draft_report.txt"
        with open(draft_path, 'w', encoding='utf-8') as f:
            f.write(content)
        
        flash('Report updated successfully')
        return redirect(url_for('preview_report'))
    
    except Exception as e:
        flash(f'Error updating report: {str(e)}')
        return redirect(url_for('preview_report'))

@app.route('/finalize_report')
def finalize_report():
    """Finalize the report, reidentify content and export to Word."""
    draft_path = OUTPUT_PATH / "draft_report.txt"
    
    if not draft_path.exists():
        flash('No draft report available. Please process your files first.')
        return redirect(url_for('index'))
    
    try:
        # Read draft report
        with open(draft_path, 'r', encoding='utf-8') as f:
            deidentified_content = f.read()
        
        # Re-identify the report
        reidentified_content = deidentifier.reidentify_text(deidentified_content)
        
        # Generate timestamp for filenames
        timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
        
        # Save the re-identified report as text
        reidentified_path = OUTPUT_PATH / f"final_report_{timestamp}.txt"
        with open(reidentified_path, 'w', encoding='utf-8') as f:
            f.write(reidentified_content)
        
        # Export to Word if supported
        if DOCX_EXPORT:
            word_path = OUTPUT_PATH / f"final_report_{timestamp}.docx"
            export_to_word(reidentified_content, word_path)
            flash('Report finalized and exported to Word format')
        else:
            flash('Report finalized (Word export not available)')
        
        return redirect(url_for('index'))
    
    except Exception as e:
        flash(f'Error finalizing report: {str(e)}')
        return redirect(url_for('preview_report'))

def export_to_word(content, output_path):
    """Export report content to Word document format."""
    doc = Document()
    
    # Set document properties
    doc.core_properties.title = "Clinical Assessment Report"
    doc.core_properties.author = "Delilah Prime"
    
    # Set styles
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    
    # Process the content line by line
    lines = content.split('\n')
    current_para = None
    
    for line in lines:
        if line.startswith('# '):
            # Heading 1
            heading = doc.add_heading(line[2:], level=1)
            heading.style.font.size = Pt(16)
            heading.style.font.bold = True
        elif line.startswith('## '):
            # Heading 2
            heading = doc.add_heading(line[3:], level=2)
            heading.style.font.size = Pt(14)
            heading.style.font.bold = True
        elif line.startswith('### '):
            # Heading 3
            heading = doc.add_heading(line[4:], level=3)
            heading.style.font.size = Pt(12)
            heading.style.font.bold = True
        elif not line.strip():
            # Empty line
            current_para = None
        else:
            # Regular paragraph
            current_para = doc.add_paragraph(line)
            
    # Add footer
    section = doc.sections[0]
    footer = section.footer
    footer_para = footer.paragraphs[0]
    footer_para.text = "Generated by Delilah Prime v1.0"
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Save document
    doc.save(output_path)
    return output_path

@app.route('/download/<filename>')
def download_file(filename):
    """Download a file from the output directory."""
    return send_from_directory(OUTPUT_PATH, filename)

@app.route('/view/<filename>')
def view_file(filename):
    """View a file from the output directory."""
    file_path = OUTPUT_PATH / filename
    
    if not file_path.exists():
        flash(f'File {filename} not found')
        return redirect(url_for('index'))
    
    with open(file_path, 'r', encoding='utf-8') as f:
        content = f.read()
    
    return render_template('view.html', filename=filename, content=content)

@app.route('/view_reference_table')
def view_reference_table():
    """View the reference table that maps placeholders to original values."""
    ref_table_path = Path.home() / Path(CONFIG["paths"]["reference_tables"])
    ref_tables = list(ref_table_path.glob("ref_table_*.json"))
    
    if not ref_tables:
        flash('No reference table found')
        return redirect(url_for('index'))
    
    # Get the most recent reference table
    latest_ref_table = max(ref_tables, key=os.path.getmtime)
    
    try:
        with open(latest_ref_table, 'r') as f:
            ref_table = json.load(f)
        
        # Organize by PHI type
        organized_table = {}
        for placeholder, value in ref_table.items():
            phi_type = placeholder.split('_')[0][1:]  # Extract type from [TYPE_id]
            if phi_type not in organized_table:
                organized_table[phi_type] = []
            organized_table[phi_type].append((placeholder, value))
        
        return render_template('reference_table.html', 
                              ref_table=organized_table, 
                              table_file=latest_ref_table.name)
    except Exception as e:
        flash(f'Error loading reference table: {str(e)}')
        return redirect(url_for('index'))

@app.route('/prompt_lab')
def prompt_lab():
    """Admin interface for managing custom prompts."""
    if not ADMIN_MODE:
        flash('Admin mode not enabled')
        return redirect(url_for('index'))
    
    prompts = []
    for file_path in PROMPTS_PATH.glob('*.txt'):
        with open(file_path, 'r', encoding='utf-8') as f:
            content = f.read()
        prompts.append({
            'name': file_path.stem,
            'content': content
        })
    
    return render_template('prompt_lab.html', prompts=prompts)

@app.route('/save_prompt', methods=['POST'])
def save_prompt():
    """Save a custom prompt template."""
    if not ADMIN_MODE:
        flash('Admin mode not enabled')
        return redirect(url_for('index'))
    
    name = request.form.get('name')
    content = request.form.get('content')
    
    if not name or not content:
        flash('Name and content required')
        return redirect(url_for('prompt_lab'))
    
    # Sanitize filename
    filename = secure_filename(f"{name}.txt")
    
    # Save prompt
    prompt_path = PROMPTS_PATH / filename
    with open(prompt_path, 'w', encoding='utf-8') as f:
        f.write(content)
    
    flash(f'Prompt template "{name}" saved successfully')
    return redirect(url_for('prompt_lab'))

@app.route('/delete_prompt/<name>')
def delete_prompt(name):
    """Delete a custom prompt template."""
    if not ADMIN_MODE:
        flash('Admin mode not enabled')
        return redirect(url_for('index'))
    
    # Sanitize filename
    filename = secure_filename(f"{name}.txt")
    
    # Delete prompt
    prompt_path = PROMPTS_PATH / filename
    if prompt_path.exists():
        os.remove(prompt_path)
        flash(f'Prompt template "{name}" deleted successfully')
    else:
        flash(f'Prompt template "{name}" not found')
    
    return redirect(url_for('prompt_lab'))

@app.route('/clear_input')
def clear_input():
    """Clear all files from the input directory."""
    for filename in os.listdir(INPUT_PATH):
        file_path = os.path.join(INPUT_PATH, filename)
        try:
            if os.path.isfile(file_path):
                os.unlink(file_path)
        except Exception as e:
            flash(f'Error deleting {filename}: {e}')
    
    flash('Input files cleared')
    log_event("All input files cleared")
    return redirect(url_for('index'))

@app.route('/clear_output')
def clear_output():
    """Clear all files from the output directory."""
    for filename in os.listdir(OUTPUT_PATH):
        file_path = os.path.join(OUTPUT_PATH, filename)
        try:
            if os.path.isfile(file_path):
                os.unlink(file_path)
        except Exception as e:
            flash(f'Error deleting {filename}: {e}')
    
    flash('Output files cleared')
    log_event("All output files cleared")
    return redirect(url_for('index'))

@app.route('/toggle_offline_mode')
def toggle_offline_mode():
    """Toggle offline mode (skip API calls)"""
    global OFFLINE_MODE
    OFFLINE_MODE = not OFFLINE_MODE
    
    if OFFLINE_MODE:
        flash('Offline mode enabled. Claude API calls will be skipped.')
    else:
        flash('Offline mode disabled. Claude API will be used to enhance content.')
    
    return redirect(url_for('index'))

@app.route('/template_manager')
def template_manager():
    """Interface for managing report templates."""
    # Get existing custom templates
    templates = []
    for file_path in CUSTOM_TEMPLATES_PATH.glob('*.json'):
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                template_data = json.load(f)
                templates.append({
                    'name': file_path.stem,
                    'description': template_data.get('description', 'No description provided'),
                    'created': template_data.get('created_date', 'Unknown'),
                    'sections': len(template_data.get('sections', {}))
                })
        except Exception as e:
            print(f"Error loading template {file_path.name}: {str(e)}")
    
    # Get section names from config
    all_sections = CONFIG['report_sections']
    
    # Load default template examples
    default_templates = {}
    for section in all_sections:
        template_path = TEMPLATES_PATH / f"{section}.txt"
        if template_path.exists():
            with open(template_path, 'r', encoding='utf-8') as f:
                default_templates[section] = f.read()
        else:
            default_templates[section] = f"# {section.replace('_', ' ').title()}\n\n{{content}}\n\n"
    
    return render_template('template_manager.html',
                          templates=templates,
                          sections=all_sections,
                          default_templates=default_templates)

@app.route('/create_template', methods=['POST'])
def create_template():
    """Create or update a template."""
    template_name = request.form.get('template_name')
    template_description = request.form.get('template_description', '')
    
    if not template_name:
        flash('Template name is required')
        return redirect(url_for('template_manager'))
    
    # Sanitize template name
    template_name = secure_filename(template_name).replace('.json', '')
    
    # Create template structure
    template_data = {
        'name': template_name,
        'description': template_description,
        'created_date': datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S'),
        'report_header': request.form.get('report_header', ''),
        'report_footer': request.form.get('report_footer', ''),
        'sections': {}
    }
    
    # Get all section data
    for section in CONFIG['report_sections']:
        section_data = {
            'prepend_content': request.form.get(f'{section}_prepend', ''),
            'append_content': request.form.get(f'{section}_append', ''),
            'default_content': request.form.get(f'{section}_default', '')
        }
        # Only add non-empty sections
        if any(section_data.values()):
            template_data['sections'][section] = section_data
    
    # Save template
    template_path = CUSTOM_TEMPLATES_PATH / f"{template_name}.json"
    with open(template_path, 'w', encoding='utf-8') as f:
        json.dump(template_data, f, indent=2)
    
    flash(f'Template "{template_name}" saved successfully')
    return redirect(url_for('template_manager'))

@app.route('/edit_template/<name>')
def edit_template(name):
    """Edit an existing template."""
    template_path = CUSTOM_TEMPLATES_PATH / f"{name}.json"
    
    if not template_path.exists():
        flash(f'Template "{name}" not found')
        return redirect(url_for('template_manager'))
    
    try:
        with open(template_path, 'r', encoding='utf-8') as f:
            template_data = json.load(f)
    except Exception as e:
        flash(f'Error loading template: {str(e)}')
        return redirect(url_for('template_manager'))
    
    # Get section names from config
    all_sections = CONFIG['report_sections']
    
    # Load default template examples
    default_templates = {}
    for section in all_sections:
        section_template_path = TEMPLATES_PATH / f"{section}.txt"
        if section_template_path.exists():
            with open(section_template_path, 'r', encoding='utf-8') as f:
                default_templates[section] = f.read()
        else:
            default_templates[section] = f"# {section.replace('_', ' ').title()}\n\n{{content}}\n\n"
    
    return render_template('edit_template.html',
                          template=template_data,
                          sections=all_sections,
                          default_templates=default_templates)

@app.route('/delete_template/<name>')
def delete_template(name):
    """Delete a template."""
    template_path = CUSTOM_TEMPLATES_PATH / f"{name}.json"
    
    if template_path.exists():
        os.remove(template_path)
        flash(f'Template "{name}" deleted successfully')
    else:
        flash(f'Template "{name}" not found')
    
    return redirect(url_for('template_manager'))

@app.route('/event_stream')
def event_stream():
    """Stream events for the activity console"""
    def generate():
        """Generate events for SSE"""
        # Send initial status message
        yield f"data: System ready - {'API Available' if claude_api.is_available() else 'API NOT Available (offline mode)'}\n\n"
        
        while True:
            try:
                # Try to get a message from the queue with a timeout
                message = event_queue.get(timeout=0.5)
                yield f"data: {message}\n\n"
                event_queue.task_done()
            except Empty:
                # If nothing in the queue, send a heartbeat comment to keep connection alive
                yield ": heartbeat\n\n"
                time.sleep(0.5)
                
    return Response(generate(), mimetype='text/event-stream')

# Monkey patch the Claude API client to log events
original_generate_narrative = claude_api.generate_narrative

def generate_narrative_with_logging(self, section_name, content):
    """Wrapper to add logging to the generate_narrative method"""
    formatted_section = section_name.replace('_', ' ').title()
    log_event(f"Starting processing of {formatted_section}")
    result = original_generate_narrative(self, section_name, content)
    log_event(f"Completed processing of {formatted_section}")
    return result

claude_api.generate_narrative = generate_narrative_with_logging.__get__(claude_api, type(claude_api))

# Monkey patch ClaudeAPIClient._apply_rate_limiting to log rate limit events
original_apply_rate_limiting = claude_api._apply_rate_limiting

def apply_rate_limiting_with_logging(self):
    """Wrapper to add logging to the _apply_rate_limiting method"""
    current_time = time.time()
    elapsed = current_time - self.last_request_time
    
    if elapsed < self.min_request_interval:
        wait_time = max(0, self.min_request_interval - elapsed)
        if wait_time > 0.1:  # Only log if waiting more than 0.1 seconds
            log_event(f"API rate limiting: Waiting {wait_time:.1f} seconds")
            time.sleep(wait_time)
    
claude_api._apply_rate_limiting = apply_rate_limiting_with_logging.__get__(claude_api, type(claude_api))

# Monkey patch the _call_claude_api method to log API requests
original_call_claude_api = claude_api._call_claude_api

def call_claude_api_with_logging(self, prompt, section_identifier):
    """Wrapper to add logging to the _call_claude_api method"""
    log_event(f"Sending API request for {section_identifier}")
    result = original_call_claude_api(self, prompt, section_identifier)
    if result:
        log_event(f"API request successful for {section_identifier}")
    else:
        log_event(f"API request failed for {section_identifier}")
    return result

claude_api._call_claude_api = call_claude_api_with_logging.__get__(claude_api, type(claude_api))

if __name__ == '__main__':
    app.run(debug=False, port=5000) 