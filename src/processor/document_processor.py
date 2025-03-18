"""
Delilah Prime - Document Processor Module

This module handles the extraction and processing of content from different document types.
"""

import os
from pathlib import Path
import re
import json
import math
from src.deidentifier.deidentifier import Deidentifier
from src.api.claude_api import ClaudeAPIClient

# Document format handlers
try:
    import docx
    from PyPDF2 import PdfReader
    DOCX_SUPPORT = True
    PDF_SUPPORT = True
except ImportError:
    DOCX_SUPPORT = False
    PDF_SUPPORT = False


class DocumentProcessor:
    """Processes different document types and extracts their content."""
    
    def __init__(self, config_path=None):
        """
        Initialize the document processor.
        
        Args:
            config_path: Path to configuration file
        """
        # Load configuration
        if config_path is None:
            config_path = Path(__file__).parent.parent.parent / "config.json"
        
        with open(config_path, 'r') as f:
            config = json.load(f)
        
        # Set up dependencies
        self.deidentifier = Deidentifier()
        self.api_client = ClaudeAPIClient()
        
        # Load section keywords
        self.section_keywords = config['section_keywords']
        
        # List of available sections from the section keywords
        self.sections = list(self.section_keywords.keys())
        
        # Document categorization patterns
        self.doc_categories = {
            'assessment_notes': [
                r'assessment', r'evaluation', r'notes', r'IHA', r'IHCAT', r'IHACAT'
            ],
            'file_review': [
                r'file.?review', r'medical.?records', r'records.?review'
            ],
            'medical_records': [
                r'medical.?records', r'chart', r'physician', r'doctor'
            ],
            'general': []  # Fallback category
        }
        
        # Configure chunking parameters
        self.max_chunk_size = config.get('max_chunk_size', 2000)  # Reduced from 4000 to 2000 characters per chunk
        self.chunk_overlap = config.get('chunk_overlap', 250)  # Reduced overlap to preserve context but avoid duplication
    
    def process_file(self, file_path):
        """
        Process a file and extract its content based on file type.
        
        Args:
            file_path: Path to the file to process
            
        Returns:
            The extracted text content
        """
        file_path = Path(file_path)
        extension = file_path.suffix.lower()
        
        if extension == '.txt' or extension == '.md':
            return self._process_text_file(file_path)
        elif extension == '.docx' and DOCX_SUPPORT:
            return self._process_docx_file(file_path)
        elif extension == '.pdf' and PDF_SUPPORT:
            return self._process_pdf_file(file_path)
        else:
            raise ValueError(f"Unsupported file type: {extension}")
    
    def split_into_chunks(self, text):
        """
        Split large text into manageable chunks to avoid API rate limits.
        Preserves paragraph breaks and includes overlap between chunks.
        
        Args:
            text: The text content to split
            
        Returns:
            List of text chunks
        """
        if len(text) <= self.max_chunk_size:
            return [text]
        
        # Split text by paragraphs to maintain context
        paragraphs = text.split('\n')
        chunks = []
        current_chunk = []
        current_length = 0
        
        for paragraph in paragraphs:
            # If adding this paragraph would exceed max size, finalize current chunk
            if current_length + len(paragraph) > self.max_chunk_size and current_chunk:
                chunks.append('\n'.join(current_chunk))
                
                # Start new chunk with overlap (keep some paragraphs from previous chunk)
                overlap_size = 0
                overlap_paragraphs = []
                
                # Add paragraphs from end of previous chunk until we reach desired overlap
                for p in reversed(current_chunk):
                    if overlap_size + len(p) <= self.chunk_overlap:
                        overlap_paragraphs.insert(0, p)
                        overlap_size += len(p) + 1  # +1 for newline
                    else:
                        break
                
                current_chunk = overlap_paragraphs
                current_length = overlap_size
            
            current_chunk.append(paragraph)
            current_length += len(paragraph) + 1  # +1 for newline
        
        # Add the last chunk if not empty
        if current_chunk:
            chunks.append('\n'.join(current_chunk))
        
        return chunks
    
    def _process_text_file(self, file_path):
        """Extract content from a text file."""
        with open(file_path, 'r', encoding='utf-8', errors='replace') as f:
            return f.read()
    
    def _process_docx_file(self, file_path):
        """Extract content from a Word document."""
        doc = docx.Document(file_path)
        
        # Extract text from paragraphs
        paragraphs = [para.text for para in doc.paragraphs]
        
        # Extract text from tables
        tables_text = []
        for table in doc.tables:
            for row in table.rows:
                row_text = ' | '.join(cell.text for cell in row.cells if cell.text.strip())
                if row_text:
                    tables_text.append(row_text)
        
        # Combine paragraphs and tables
        all_text = paragraphs + tables_text
        return '\n'.join(all_text)
    
    def _process_pdf_file(self, file_path):
        """Extract content from a PDF file."""
        reader = PdfReader(file_path)
        text_parts = []
        
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text)
        
        # Join all pages with a separator
        return '\n\n'.join(text_parts)
    
    def categorize_document(self, file_name):
        """
        Determine the document type based on filename patterns.
        
        Args:
            file_name: Name of the file
            
        Returns:
            Document type category
        """
        file_name = file_name.lower()
        
        # Check for assessment notes
        if any(term in file_name for term in ['assessment', 'eval', 'notes', 'report', 'ot report', 'pt report']):
            return 'assessment_notes'
        # Check for file reviews
        elif any(term in file_name for term in ['file', 'review', 'history', 'record', 'documentation']):
            return 'file_review'
        # Check for medical documents
        elif any(term in file_name for term in ['medical', 'health', 'clinical', 'hospital', 'discharge']):
            return 'medical_documents'
        else:
            return 'other'
    
    def organize_content(self, documents):
        """
        Organize document content by sections based on keywords.
        
        Args:
            documents: List of file paths or dictionary of document content 
            
        Returns:
            Dictionary with section content
        """
        organized_content = {section: "" for section in self.section_keywords.keys()}
        
        # Convert list of file paths to dict of content if needed
        documents_dict = {}
        if isinstance(documents, list):
            # Process each file to get content
            for file_path in documents:
                try:
                    file_content = self._extract_content_from_file(file_path)
                    doc_name = os.path.basename(file_path)
                    documents_dict[doc_name] = file_content
                    print(f"Extracted content from {doc_name}")
                except Exception as e:
                    print(f"Error extracting content from {file_path}: {str(e)}")
        else:
            # Already a dictionary
            documents_dict = documents
            
        # Now process the dictionary of content
        for doc_name, content in documents_dict.items():
            # Skip empty content
            if not content:
                continue
                
            print(f"Organizing content from: {doc_name}")
            
            # Try to categorize the document type
            doc_type = self.categorize_document(doc_name)
            
            # Process each section
            for section_name, keywords in self.section_keywords.items():
                # Extract content based on keywords
                section_content = self._extract_section(content, keywords)
                
                # If we found content, add it to the organized content
                if section_content:
                    if organized_content[section_name]:
                        organized_content[section_name] += f"\n\n--- From {doc_name} ---\n\n"
                    else:
                        organized_content[section_name] += f"--- From {doc_name} ---\n\n"
                        
                    organized_content[section_name] += section_content
                    
        return organized_content
    
    def _extract_section(self, content, keywords, context_lines=15):
        """
        Extract a section of text based on keywords.
        
        Args:
            content: Text content to extract from
            keywords: List of keywords that might indicate the section
            context_lines: Number of lines to include for context
            
        Returns:
            Extracted section text
        """
        # Convert content to lowercase for case-insensitive matching
        content_lower = content.lower()
        lines = content.split('\n')
        
        # First try to find section headers
        for i, line in enumerate(lines):
            line_lower = line.lower()
            
            # Check if line contains a section header with any of the keywords
            if any(keyword in line_lower for keyword in keywords):
                # Check if this looks like a header (short line, possibly with formatting)
                if len(line.strip()) < 100 and (line.strip().endswith(':') or line.strip().startswith('#') or line.isupper()):
                    start = max(0, i)
                    
                    # Look for the next section header or end of content
                    end = len(lines)
                    for j in range(i + 1, min(len(lines), i + 50)):
                        next_line = lines[j].lower()
                        # Check if this line looks like a new section header
                        if (len(next_line.strip()) < 100 and 
                            (next_line.strip().endswith(':') or next_line.strip().startswith('#') or next_line.isupper()) and
                            any(kw in next_line for kw in sum(self.section_keywords.values(), []))):
                            end = j
                            break
                    
                    return '\n'.join(lines[start:end])
        
        # If no clear section headers, try to find paragraphs containing keywords
        best_match = None
        best_score = 0
        
        for i, line in enumerate(lines):
            line_lower = line.lower()
            
            # Calculate a score based on how many keywords are found
            score = sum(2 for keyword in keywords if keyword in line_lower)
            
            # Check surrounding lines for context
            context_score = 0
            for j in range(max(0, i-5), min(len(lines), i+5)):
                if i != j:  # Skip the current line
                    context_score += sum(1 for keyword in keywords if keyword in lines[j].lower())
            
            total_score = score + context_score
            if total_score > best_score:
                best_score = total_score
                best_match = i
        
        if best_match is not None and best_score >= 2:
            start = max(0, best_match - 5)
            end = min(len(lines), best_match + context_lines)
            return '\n'.join(lines[start:end])
        
        # If no section found, return empty string
        return ""

    def process_documents(self, files, prompt_template=None, reference_table=None):
        """
        Process multiple document files, enhancing the content if API is available.
        
        Args:
            files: List of file paths to process
            prompt_template: Custom prompt template to use for enhancement (optional)
            reference_table: Reference table for re-identification (optional)
            
        Returns:
            report_content: Dictionary with section content
        """
        print("\n== STARTING DOCUMENT PROCESSING ==")
        print(f"Processing {len(files)} files")
        
        # Use the organize_content method which already exists
        organized_content = self.organize_content(files)
        
        # Create a dictionary to store the final report content
        report_content = {}
        
        # Process each section
        for section_name, content in organized_content.items():
            # Skip if this section is empty
            if not content:
                print(f"Skipping empty section: {section_name}")
                continue
                
            print(f"Processing section: {section_name}")
            
            # Add content to the report (will be enhanced if API available)
            report_content[section_name] = content
            
            # If the API is available, attempt enhancement
            if self.api_client.is_available():
                try:
                    # If custom prompt is provided, use it
                    if prompt_template:
                        enhanced_content = self.api_client.generate_custom_narrative(
                            section_name, content, prompt_template
                        )
                    else:
                        # Use the default prompt
                        enhanced_content = self.api_client.generate_narrative(
                            section_name, content
                        )
                    
                    # Only update if we got back valid content
                    if enhanced_content and enhanced_content.strip():
                        report_content[section_name] = enhanced_content
                except Exception as e:
                    print(f"Error enhancing section {section_name}: {str(e)}")
                    # Keep the original content on error
                    
        print("== DOCUMENT PROCESSING COMPLETE ==")
        
        # Perform re-identification if reference table is provided
        if reference_table:
            for section in report_content:
                report_content[section] = self.deidentifier.reidentify_content(
                    report_content[section], reference_table
                )
                
        return report_content

    def _extract_content_from_file(self, file_path):
        """
        Extract text content from a file based on its type.
        
        Args:
            file_path: Path to the file
            
        Returns:
            De-identified content from the file
        """
        file_ext = os.path.splitext(file_path)[1].lower()
        
        # Extract content based on file type
        if file_ext == '.txt':
            content = self._process_text_file(file_path)
        elif file_ext == '.docx':
            content = self._process_docx_file(file_path)
        elif file_ext == '.pdf':
            content = self._process_pdf_file(file_path)
        else:
            print(f"Unsupported file type: {file_ext}")
            return ""
            
        # De-identify content
        if content:
            return self.deidentifier.deidentify_text(content)
        
        return ""
        
    def _organize_content_by_section(self, content, organized_content=None):
        """
        Organize content into sections based on keywords.
        
        Args:
            content: The document content to organize
            organized_content: Dictionary to store organized content, if None a new one is created
            
        Returns:
            Dictionary with content organized by section
        """
        if organized_content is None:
            organized_content = {section: "" for section in self.section_keywords.keys()}
            
        # Skip if no content
        if not content:
            return organized_content
            
        # Process each section
        for section_name, keywords in self.section_keywords.items():
            # Extract content for this section based on keywords
            section_content = self._extract_section(content, keywords)
            
            # If we found content, add it to the organized content
            if section_content:
                if organized_content[section_name]:
                    organized_content[section_name] += f"\n\n{section_content}"
                else:
                    organized_content[section_name] = section_content
                    
        return organized_content 